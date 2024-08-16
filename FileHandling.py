from PyQt5.QtWidgets import (QTableWidget, QLineEdit)
import datetime
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import shutil
import re
import sys

class FileManipulation():
    def __init__(self, table_widget: QTableWidget, two_week_table: QTableWidget, name_input: QLineEdit):
        self.table_widget = table_widget
        self.two_week_table = two_week_table
        self.name_input = name_input
       
    def notOnlyLetters(self, name):
        # Return True if anything other than characters and spaces
        return bool(re.search(r'[^a-zA-Z ]', name)) #The ^ inside the brackets negates the character class, meaning the pattern matches any character except letters and spaces.

        
    def generate_docx(self, filename):
        today = datetime.date.today()
        doc = Document()
        
        # Add title
        doc.add_heading('Rest Period Acknowledgment Form', 0)
        
        # Add formatted content from information1.txt
        info1_content = self.read_and_format_file('information1.txt')
        self.add_html_content(doc, info1_content)
        
        # Add Table Data
        doc.add_paragraph('Table Data:')
        self.add_table(doc, self.table_widget)
        
        doc.add_paragraph("")
        
        # Add formatted content from information2.txt
        info2_content = self.read_and_format_file('information2.txt')
        self.add_html_content(doc, info2_content)
        
        # Add Two Week Table Data
        doc.add_paragraph('Two Week Table:')
        self.add_table(doc, self.two_week_table)
        
        # Add Name with formatting
        doc.add_paragraph("")
        doc.add_paragraph(f"Name: {self.name_input.text()}")
        
        # Add Signature and date
        doc.add_paragraph(f"Signature: ______________________________________________________                       Date: {today}")
        
        # Save the document
        doc.save(filename)
    
    def add_html_content(self, doc, html_content):
        # Use BeautifulSoup to parse HTML content
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Convert HTML content to Word paragraphs
        for element in soup:
            if element.name == 'p':  # Paragraphs
                doc.add_paragraph(element.get_text())
            elif element.name == 'br':  # Line breaks
                doc.add_paragraph()  # Add an empty paragraph for line break
            
    def add_table(self, doc, table_widget):
        num_rows = table_widget.rowCount()
        num_cols = table_widget.columnCount()
        
        # Check if the table has headers (e.g., days of the week)
        has_headers = num_rows == 3
        
        # Create the table
        if has_headers:
            # Create table with an extra row for headers
            table = doc.add_table(rows=num_rows + 1, cols=num_cols)
            header_row = table.rows[0]  # Header row
            
            # Add headers
            headers = ['Sunday', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
            for col_idx, header in enumerate(headers):
                cell = header_row.cells[col_idx]
                cell.text = header
                # Set font size for header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
                cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            # Fill the rest of the table with data
            for row in range(1, num_rows + 1):  # Skip the header row
                for col in range(num_cols):
                    cell = table.cell(row, col)
                    item = table_widget.item(row - 1, col)  # Adjust index for header
                    cell.text = item.text() if item else ''
                    # Set font size and alignment
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
        else:
            # Create table without headers
            table = doc.add_table(rows=num_rows, cols=num_cols)
            
            # Fill the table with data
            for row in range(num_rows):
                for col in range(num_cols):
                    cell = table.cell(row, col)
                    item = table_widget.item(row, col)
                    cell.text = item.text() if item else ''
                    # Set font size and alignment
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(11)
                        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set column widths if necessary
        for col in table.columns:
            for cell in col.cells:
                cell.width = Inches(1.0)
        
        # Set uniform row height
        uniform_row_height = Pt(14)
        for row in table.rows:
            for cell in row.cells:
                # Set the height of each cell
                cell.height = uniform_row_height
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
        
        # Optional: Set table formatting
        table.style = 'Table Grid'
        
        # Add an empty paragraph to ensure no extra spacing before the table
        doc.add_paragraph()

        
    def delete_doc(self, filename):
        # grab doc file path of doc
        directory = os.path.dirname(os.path.abspath(__file__))
        file_path = os.path.join(directory, filename)
        
        # deletes doc file if it exists
        if os.path.isfile(file_path):
            os.remove(file_path)
            print(f"File '{filename}' has been deleted.")
        else:
            print(f"File '{filename}' does not exist.")
    
    def move_pdf(self, filename):
        # Get the source path (same directory as main.py)
        source_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), filename)

        # Determine the base directory depending on operating system
        if os.name == 'nt':  # Windows
            base_directory = os.environ.get('USERPROFILE', '')  # This gets 'C:\Users\Username'
        elif os.name == 'posix':  # macOS/Linux
            base_directory = os.environ.get('HOME', '')  # This gets '/home/username'
        else:
            raise EnvironmentError('Unsupported operating system')
        
        start_date = self.two_week_table.item(0, 0).text() # grab the start date from the table
        end_date = self.two_week_table.item(1, 6).text() # grab the end date from the table

        # Define the rest of the path relative to the base directory
        dates = f"PayPeriod_{start_date}_{end_date}" 
        #relative_path = f'Special Services Group, LLC\SSG Customer Access - Documents\Babisha Mudaliar\Rest Period Acknowledgment Form\{dates}'
        relative_path = os.path.join('Special Services Group, LLC', 'SSG Customer Access - Documents', 'Babisha Mudaliar', 'Rest Period Acknowledgment Form' ,dates)

        # Construct the full destination path
        destination_directory1 = os.path.join(base_directory, relative_path)
        destination_path1 = os.path.join(destination_directory1, filename)
        destination_directory2 = os.path.join("D:", '\Special Services Group, LLC', 'SSG Customer Access - Documents', 'Babisha Mudaliar', 'Rest Period Acknowledgment Form', dates)
        destination_path2 = os.path.join(destination_directory2, filename)
    
        destination1_exists = os.path.join(base_directory, 'Special Services Group, LLC', 'SSG Customer Access - Documents', 'Babisha Mudaliar')
        destination2_exists = os.path.join('D:', '\Special Services Group, LLC', 'SSG Customer Access - Documents', 'Babisha Mudaliar')
    
        # Check if the source file exists
        if os.path.isfile(source_path):
            # Handle the first destination
            if os.path.exists(destination1_exists):
                if not os.path.exists(destination_directory1):
                    print(f"Creating directory: {destination_directory1}")
                    os.makedirs(destination_directory1)  # Create the directory if it does not exist
    
                try:
                    shutil.move(source_path, destination_path1)
                    print(f"File successfully moved to {destination_path1}")
                except FileNotFoundError:
                    print("Failed to move the file to destination1.")
                except Exception as e:
                    print(f"Error moving file to destination1: {e}")
            else:
                print(f"Base path for destination1 does not exist: {destination1_exists}")
    
            # Handle the second destination
            if os.path.exists(destination2_exists):
                if not os.path.exists(destination_directory2):
                    print(f"Creating directory: {destination_directory2}")
                    os.makedirs(destination_directory2)  # Create the directory if it does not exist
    
                try:
                    shutil.move(source_path, destination_path2)
                    print(f"File successfully moved to {destination_path2}")
                except FileNotFoundError:
                    print("Failed to move the file to destination2.")
                except Exception as e:
                    print(f"Error moving file to destination2: {e}")
            else:
                print(f"Base path for destination2 does not exist: {destination2_exists}")
        else:
            print("Source file does not exist.")
        
    def read_and_format_file(self, filename):
        try:
            with open(filename, 'r', encoding='utf-8') as file:
                content = file.read()
                # Directly return content if it's in proper HTML format
                return content
        except FileNotFoundError:
            return f"Error: The file '{filename}' was not found."
        except Exception as e:
            return f"An error occurred: {e}"
    
        